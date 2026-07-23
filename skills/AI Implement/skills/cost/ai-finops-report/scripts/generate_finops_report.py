#!/usr/bin/env python3
"""AI FinOps report generator for M365 Copilot Cowork / Neptune credit exports.

Reads the credit-consumption CSV exports (per-user / per-group / per-service /
Cowork task details), computes FinOps insights with a cost-optimization focus,
and writes:
  * one Word narrative (.docx)
  * one Excel workbook (.xlsx)

Files are auto-detected by header signature, so filenames don't matter.

Usage:
    python3 generate_finops_report.py \
        --input-dir "/path/to/csvs" \
        --output-dir "/path/to/out" \
        --period "2026-07" \
        --org "Contoso" \
        [--credit-cost 0.01 --currency USD]

Dependencies: python-docx, openpyxl
    pip install python-docx openpyxl
"""

from __future__ import annotations

import argparse
import csv
import glob
import os
import statistics
import sys
from datetime import datetime

# ---------------------------------------------------------------------------
# Tunable thresholds (see references/finops-playbook.md)
# ---------------------------------------------------------------------------
CONCENTRATION_FLAG = 0.60      # one user >= 60% of total spend -> risk
NEAR_LIMIT = 0.80              # % used >= 80% -> near limit
OVER_LIMIT = 1.00             # % used > 100% -> breach
OVER_CAP_HEADROOM = 0.40      # limit set but < 40% used -> over-provisioned cap
LOW_AUTOMATION = 0.20         # scheduled/total < 20% -> low automation maturity
RIGHTSIZE_MULTIPLIER = 1.2    # suggested new limit = 1.2x observed usage


# ---------------------------------------------------------------------------
# Parsing helpers
# ---------------------------------------------------------------------------
def to_num(value) -> float:
    """Parse a numeric cell, tolerating thousands separators and blanks."""
    if value is None:
        return 0.0
    s = str(value).strip().replace(",", "").replace("%", "")
    if s == "" or s.lower() in {"n/a", "na", "none", "null"}:
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def fmt_num(value) -> str:
    """Integer-looking numbers without decimals, else 1 dp."""
    n = float(value)
    if abs(n - round(n)) < 1e-9:
        return f"{int(round(n)):,}"
    return f"{n:,.1f}"


def fmt_pct(value) -> str:
    return f"{value * 100:.1f}%"


def parse_date(value) -> str:
    if not value:
        return ""
    s = str(value).strip()
    if not s:
        return ""
    for fmt in ("%Y-%m-%dT%H:%M:%S.%f", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d %H:%M:%S",
                "%Y-%m-%d", "%m/%d/%Y %I:%M:%S %p", "%m/%d/%Y"):
        try:
            return datetime.strptime(s, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return s[:10]


def read_csv(path: str) -> list[dict]:
    with open(path, "r", encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


def headers_of(path: str) -> list[str]:
    with open(path, "r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.reader(fh)
        try:
            return next(reader)
        except StopIteration:
            return []


# ---------------------------------------------------------------------------
# Detection + ingestion
# ---------------------------------------------------------------------------
def detect_type(headers: list[str]) -> str | None:
    hs = {h.strip() for h in headers}
    if "User Principal Name" in hs and "Monthly credit limit" in hs:
        return "users"
    if "Group ID" in hs:
        return "groups"
    if "Service Name" in hs:
        return "services"
    if "TotalTasks" in hs:
        return "tasks"
    return None


def load_csvs(input_dir: str) -> dict:
    data = {"users": [], "groups": [], "services": [], "tasks": []}
    found = {}
    for path in sorted(glob.glob(os.path.join(input_dir, "*.csv"))):
        kind = detect_type(headers_of(path))
        if kind is None:
            print(f"  ! skip (unrecognized headers): {os.path.basename(path)}")
            continue
        rows = read_csv(path)
        data[kind].extend(rows)
        found.setdefault(kind, []).append(os.path.basename(path))
    return data, found


# ---------------------------------------------------------------------------
# Insight computation
# ---------------------------------------------------------------------------
def compute_insights(data: dict, credit_cost: float | None) -> dict:
    users = []
    for r in data["users"]:
        used = to_num(r.get("Monthly credits used"))
        limit = to_num(r.get("Monthly credit limit"))
        sessions = to_num(r.get("Session Count"))
        licensed = str(r.get("Microsoft 365 Copilot license", "")).strip().lower() == "yes"
        pct = (used / limit) if limit > 0 else None
        users.append({
            "name": (r.get("Display Name") or "").strip(),
            "upn": (r.get("User Principal Name") or "").strip(),
            "limit": limit,
            "used": used,
            "sessions": sessions,
            "licensed": licensed,
            "last_active": parse_date(r.get("Last activity date")),
            "pct": pct,
            "per_session": (used / sessions) if sessions > 0 else 0.0,
        })

    groups = []
    for r in data["groups"]:
        gid = (r.get("Group ID") or "").strip()
        groups.append({
            "name": (r.get("Display Name") or "").strip() or "(未分組 / Ungrouped)",
            "total_users": to_num(r.get("Total Users")),
            "used": to_num(r.get("Monthly credits used")),
            "active_members": to_num(r.get("Members that used credits")),
            "avg_per_user_day": to_num(r.get("Avg credits per user per day")),
            "gid": gid,
            "sessions": to_num(r.get("Session Count")),
            "last_active": parse_date(r.get("Last activity date")),
        })

    services = []
    for r in data["services"]:
        services.append({
            "name": (r.get("Service Name") or "").strip(),
            "active_users": to_num(r.get("Active users")),
            "used": to_num(r.get("Monthly credits used")),
            "last_active": parse_date(r.get("Last activity date")),
        })

    tasks = []
    for r in data["tasks"]:
        total = to_num(r.get("TotalTasks"))
        sched = to_num(r.get("ScheduledTasks"))
        tasks.append({
            "name": (r.get("DisplayName") or "").strip(),
            "upn": (r.get("UserPrincipalName") or "").strip(),
            "total": total,
            "scheduled": sched,
            "user_initiated": to_num(r.get("UserInitiatedTasks")),
            "active_days": to_num(r.get("ActiveDays")),
            "last_active": parse_date(r.get("LastActivityDate")),
            "automation_ratio": (sched / total) if total > 0 else None,
        })

    total_used = sum(u["used"] for u in users)
    active_users = [u for u in users if u["used"] > 0]
    top = sorted(users, key=lambda u: u["used"], reverse=True)

    concentration = (top[0]["used"] / total_used) if (top and total_used > 0) else 0.0

    per_session_vals = [u["per_session"] for u in active_users if u["per_session"] > 0]
    median_ps = statistics.median(per_session_vals) if per_session_vals else 0.0

    return {
        "users": users,
        "groups": groups,
        "services": services,
        "tasks": tasks,
        "total_used": total_used,
        "active_count": len(active_users),
        "user_count": len(users),
        "avg_per_active": (total_used / len(active_users)) if active_users else 0.0,
        "top": top,
        "concentration": concentration,
        "median_per_session": median_ps,
        "credit_cost": credit_cost,
    }


# ---------------------------------------------------------------------------
# Recommendation engine (the optimization payload)
# ---------------------------------------------------------------------------
def build_recommendations(ins: dict) -> list[dict]:
    recs = []
    users = ins["users"]
    total = ins["total_used"]

    # 1. Ungoverned / unlimited spenders
    unlimited = [u for u in users if u["limit"] == 0 and u["used"] > 0]
    if unlimited:
        worst = max(unlimited, key=lambda u: u["used"])
        suggested = int(worst["used"] * RIGHTSIZE_MULTIPLIER)
        recs.append({
            "priority": "High",
            "theme": "封頂未受管控的用量 / Cap ungoverned spend",
            "finding": f"{len(unlimited)} 位使用者沒有月度額度上限（limit=0），"
                       f"其中 {worst['name'] or worst['upn']} 已用 {fmt_num(worst['used'])} credits。",
            "action": f"為這些帳號設定月度額度，建議上限 ≈ 觀察用量的 {RIGHTSIZE_MULTIPLIER}×"
                      f"（例：{worst['name'] or worst['upn']} 設 ~{suggested:,} credits）。",
            "benefit": "支出可預測、避免失控成長。",
        })

    # 2. Concentration risk
    if ins["concentration"] >= CONCENTRATION_FLAG and ins["top"]:
        t = ins["top"][0]
        recs.append({
            "priority": "High",
            "theme": "集中度風險 / Concentration risk",
            "finding": f"單一使用者（{t['name'] or t['upn']}）佔全部消耗的 "
                       f"{fmt_pct(ins['concentration'])}（{fmt_num(t['used'])}/{fmt_num(total)} credits）。",
            "action": "檢視該使用者的工作負載，將重複性 / 可用授權內工具完成的任務"
                      "改走 Copilot Chat + Work IQ（見 copilot-cost-advisor）。",
            "benefit": "直接壓低最大的單一成本來源。",
        })

    # 3. Wasted licences
    wasted = [u for u in users if u["licensed"] and u["used"] == 0 and u["sessions"] == 0]
    if wasted:
        recs.append({
            "priority": "High",
            "theme": "回收閒置授權 / Reclaim wasted licences",
            "finding": f"{len(wasted)} 位持有 M365 Copilot 授權但本期 0 用量、0 session。",
            "action": "確認是否仍需要，將授權重新指派給有需求者或釋出。",
            "benefit": "直接節省授權成本。",
        })

    # 4. Over-provisioned caps
    over_cap = [u for u in users if u["pct"] is not None and 0 < u["pct"] < OVER_CAP_HEADROOM]
    if over_cap:
        recs.append({
            "priority": "Medium",
            "theme": "調降過高的額度 / Right-size over-caps",
            "finding": f"{len(over_cap)} 位使用者的實際用量低於額度的 {int(OVER_CAP_HEADROOM*100)}%。",
            "action": "將月度額度向實際用量靠攏，保留合理緩衝即可。",
            "benefit": "更精準的預算控制，不影響使用者體驗。",
        })

    # 5. Near / over limit
    near = [u for u in users if u["pct"] is not None and u["pct"] >= NEAR_LIMIT]
    if near:
        breaches = [u for u in near if u["pct"] > OVER_LIMIT]
        recs.append({
            "priority": "Medium",
            "theme": "接近 / 超出額度 / Near or over limit",
            "finding": f"{len(near)} 位使用者用量 ≥ 額度 80%"
                       + (f"，其中 {len(breaches)} 位已超出上限。" if breaches else "。"),
            "action": "確認是否為正常需求；若是，審慎調高額度，否則輔導使用方式。",
            "benefit": "避免被限流中斷或預算超支。",
        })

    # 6. Automation maturity
    tasks = [t for t in ins["tasks"] if t["automation_ratio"] is not None]
    low_auto = [t for t in tasks if t["automation_ratio"] < LOW_AUTOMATION]
    if tasks and low_auto:
        recs.append({
            "priority": "Low",
            "theme": "提升自動化成熟度 / Raise automation maturity",
            "finding": f"{len(low_auto)}/{len(tasks)} 位使用者的排程任務佔比低於 "
                       f"{int(LOW_AUTOMATION*100)}%（多為臨時手動觸發）。",
            "action": "把重複性工作轉為排程任務（Scheduled），減少臨時 run。",
            "benefit": "每個 credit 產出更多成果。",
        })

    # 7. Dormant groups
    dormant = [g for g in ins["groups"] if g["total_users"] > 0 and g["active_members"] == 0]
    if dormant:
        recs.append({
            "priority": "Low",
            "theme": "清理閒置群組 / Retire dormant groups",
            "finding": f"{len(dormant)} 個群組有成員但本期 0 位使用 credits。",
            "action": "檢視群組是否仍需要，清理或重新指派。",
            "benefit": "降低管理複雜度與授權浪費。",
        })

    order = {"High": 0, "Medium": 1, "Low": 2}
    recs.sort(key=lambda r: order.get(r["priority"], 9))
    return recs


# ---------------------------------------------------------------------------
# Excel output
# ---------------------------------------------------------------------------
def write_excel(ins: dict, recs: list[dict], out_path: str, meta: dict) -> None:
    from openpyxl import Workbook
    from openpyxl.chart import (BarChart, DoughnutChart, LineChart, PieChart,
                                Reference, ScatterChart, Series)
    from openpyxl.chart.marker import DataPoint
    from openpyxl.chart.shapes import GraphicalProperties
    from openpyxl.drawing.line import LineProperties
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
    from openpyxl.worksheet.table import Table, TableStyleInfo

    # Power BI-ish palette
    PALETTE = ["4E79A7", "F28E2B", "E15759", "76B7B2", "59A14F", "EDC948",
               "B07AA1", "FF9DA7", "9C755F", "BAB0AC", "86BCB6", "D37295"]
    ACCENT = "FFB900"
    CARD_COLORS = ["2E5A88", "1F7A54", "8A5A00", "9C3D3D", "5A3A7A", "0F6E6E"]
    INK = "1F2937"
    HEADER_FILL = PatternFill("solid", fgColor=ACCENT)
    HEADER_FONT = Font(bold=True, color="000000")
    PRIO_FILL = {
        "High": PatternFill("solid", fgColor="F8CBAD"),
        "Medium": PatternFill("solid", fgColor="FFE699"),
        "Low": PatternFill("solid", fgColor="C6E0B4"),
    }

    wb = Workbook()

    def style_header(ws, row, ncols):
        for c in range(1, ncols + 1):
            cell = ws.cell(row=row, column=c)
            cell.fill = HEADER_FILL
            cell.font = HEADER_FONT
            cell.alignment = Alignment(horizontal="center", vertical="center")

    def autosize(ws, widths):
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def make_table(ws, name):
        """Turn the used range into a native Excel Table (filter dropdowns)."""
        if ws.max_row < 2:
            return
        ref = f"A1:{get_column_letter(ws.max_column)}{ws.max_row}"
        t = Table(displayName=name, ref=ref)
        t.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2",
                                          showRowStripes=True,
                                          showColumnStripes=False)
        ws.add_table(t)

    # =====================================================================
    # Dashboard (Power BI style) + hidden _calc source sheet
    # =====================================================================
    dash = wb.active
    dash.title = "📊 Dashboard"
    dash.sheet_view.showGridLines = False
    calc = wb.create_sheet("_calc")
    calc.sheet_state = "hidden"

    for col in range(1, 26):
        dash.column_dimensions[get_column_letter(col)].width = 10.5

    # Title band
    dash.merge_cells("B1:R1")
    tcell = dash["B1"]
    tcell.value = f"AI FinOps 儀表板 — {meta['org']}"
    tcell.font = Font(bold=True, size=18, color=INK)
    tcell.alignment = Alignment(vertical="center")
    dash.row_dimensions[1].height = 28
    dash.merge_cells("B2:R2")
    scell = dash["B2"]
    scell.value = f"期間 {meta['period']}　|　產出 {meta['generated']}　|　單位：credits"
    scell.font = Font(size=10, color="6B7280")

    # KPI cards
    def card(col, label, value, fill):
        c1, c2 = get_column_letter(col), get_column_letter(col + 2)
        dash.merge_cells(f"{c1}4:{c2}4")
        dash.merge_cells(f"{c1}5:{c2}6")
        for rr in (4, 5, 6):
            for cc in range(col, col + 3):
                dash.cell(row=rr, column=cc).fill = PatternFill("solid", fgColor=fill)
        lab = dash[f"{c1}4"]
        lab.value = label
        lab.font = Font(bold=True, size=9, color="FFFFFF")
        lab.alignment = Alignment(horizontal="center", vertical="center")
        val = dash[f"{c1}5"]
        val.value = value
        val.font = Font(bold=True, size=20, color="FFFFFF")
        val.alignment = Alignment(horizontal="center", vertical="center")

    unlimited_n = sum(1 for u in ins["users"] if u["limit"] == 0 and u["used"] > 0)
    idle_n = sum(1 for u in ins["users"]
                 if u["licensed"] and u["used"] == 0 and u["sessions"] == 0)
    cards = [
        ("總消耗 Total credits", fmt_num(ins["total_used"])),
        ("活躍使用者 Active", f"{ins['active_count']}/{ins['user_count']}"),
        ("平均/使用者 Avg", fmt_num(ins["avg_per_active"])),
        ("最大佔比 Top share", fmt_pct(ins["concentration"])),
        ("無上限帳號 Unlimited", str(unlimited_n)),
        ("閒置授權 Idle lic.", str(idle_n)),
    ]
    dash.row_dimensions[4].height = 16
    dash.row_dimensions[5].height = 20
    dash.row_dimensions[6].height = 18
    for i, (lab, val) in enumerate(cards):
        card(2 + i * 3, lab, val, CARD_COLORS[i % len(CARD_COLORS)])

    def color_bar(chart, hexcolor):
        if chart.series:
            chart.series[0].graphicalProperties = GraphicalProperties(solidFill=hexcolor)

    def color_points(chart, n):
        for i in range(n):
            pt = DataPoint(idx=i)
            pt.graphicalProperties = GraphicalProperties(
                solidFill=PALETTE[i % len(PALETTE)])
            chart.series[0].data_points.append(pt)

    def section(anchor_cell, text):
        dash[anchor_cell] = text
        dash[anchor_cell].font = Font(bold=True, size=11, color=INK)

    # ---- build _calc source tables ----
    # Top spenders + cumulative %  (cols 1,2,3)
    top = [u for u in ins["top"] if u["used"] > 0][:12]
    calc["A1"] = "User"; calc["B1"] = "Credits"; calc["C1"] = "Cumulative %"
    run = 0.0
    for i, u in enumerate(top):
        run += u["used"]
        calc.cell(row=2 + i, column=1, value=u["name"] or u["upn"])
        calc.cell(row=2 + i, column=2, value=u["used"])
        calc.cell(row=2 + i, column=3,
                  value=round(run / ins["total_used"] * 100, 1) if ins["total_used"] else 0)
    ntop = len(top)

    # Service split (cols 5,6)
    calc["E1"] = "Service"; calc["F1"] = "Credits"
    svc = sorted(ins["services"], key=lambda x: x["used"], reverse=True)
    for i, s in enumerate(svc):
        calc.cell(row=2 + i, column=5, value=s["name"])
        calc.cell(row=2 + i, column=6, value=s["used"])
    nsvc = len(svc)

    # Credits by group (cols 8,9)
    calc["H1"] = "Group"; calc["I1"] = "Credits"
    grp = sorted([g for g in ins["groups"] if g["used"] > 0],
                 key=lambda x: x["used"], reverse=True)[:10]
    for i, g in enumerate(grp):
        calc.cell(row=2 + i, column=8, value=g["name"])
        calc.cell(row=2 + i, column=9, value=g["used"])
    ngrp = len(grp)

    # Utilization % for capped users (cols 11,12)
    calc["K1"] = "User"; calc["L1"] = "% Used"
    util = sorted([u for u in ins["users"] if u["pct"] is not None and u["used"] > 0],
                  key=lambda x: x["pct"], reverse=True)[:12]
    for i, u in enumerate(util):
        calc.cell(row=2 + i, column=11, value=u["name"] or u["upn"])
        calc.cell(row=2 + i, column=12, value=round(u["pct"] * 100, 1))
    nutil = len(util)

    # Licence status (cols 14,15)
    calc["N1"] = "Status"; calc["O1"] = "Users"
    lic_active = sum(1 for u in ins["users"] if u["licensed"] and u["used"] > 0)
    lic_idle = sum(1 for u in ins["users"]
                   if u["licensed"] and u["used"] == 0 and u["sessions"] == 0)
    unlicensed = sum(1 for u in ins["users"] if not u["licensed"])
    lic_rows = [("Active licensed", lic_active), ("Idle licensed", lic_idle),
                ("Unlicensed", unlicensed)]
    lic_rows = [r for r in lic_rows if r[1] > 0]
    for i, (lab, cnt) in enumerate(lic_rows):
        calc.cell(row=2 + i, column=14, value=lab)
        calc.cell(row=2 + i, column=15, value=cnt)
    nlic = len(lic_rows)

    # Sessions vs credits scatter (cols 17,18)
    calc["Q1"] = "Sessions"; calc["R1"] = "Credits"
    scatter = [u for u in ins["users"] if u["used"] > 0]
    for i, u in enumerate(scatter):
        calc.cell(row=2 + i, column=17, value=u["sessions"])
        calc.cell(row=2 + i, column=18, value=u["used"])
    nsc = len(scatter)

    # Automation split (cols 20,21,22)
    calc["T1"] = "User"; calc["U1"] = "Scheduled"; calc["V1"] = "User-initiated"
    autot = sorted([t for t in ins["tasks"] if t["total"] > 0],
                   key=lambda x: x["total"], reverse=True)[:12]
    for i, t in enumerate(autot):
        calc.cell(row=2 + i, column=20, value=t["name"] or t["upn"])
        calc.cell(row=2 + i, column=21, value=int(t["scheduled"]))
        calc.cell(row=2 + i, column=22, value=int(t["user_initiated"]))
    nauto = len(autot)

    # ---- charts on dashboard ----
    CW, CH = 15, 8

    # 1. Top spenders (horizontal bar)
    if ntop:
        section("B8", "🏆 Top Spenders (credits)")
        ch = BarChart(); ch.type = "bar"; ch.title = None; ch.legend = None
        ch.width, ch.height = CW, CH
        ch.add_data(Reference(calc, min_col=2, min_row=1, max_row=ntop + 1),
                    titles_from_data=True)
        ch.set_categories(Reference(calc, min_col=1, min_row=2, max_row=ntop + 1))
        color_bar(ch, PALETTE[0])
        dash.add_chart(ch, "B9")

    # 2. Service split (doughnut)
    if nsvc:
        section("K8", "🧩 Service split")
        ch = DoughnutChart(); ch.title = None
        ch.width, ch.height = CW, CH
        ch.add_data(Reference(calc, min_col=6, min_row=1, max_row=nsvc + 1),
                    titles_from_data=True)
        ch.set_categories(Reference(calc, min_col=5, min_row=2, max_row=nsvc + 1))
        color_points(ch, nsvc)
        dash.add_chart(ch, "K9")

    # 3. Concentration Pareto (bar + cumulative line)
    if ntop:
        section("B25", "📈 Concentration (Pareto)")
        bar = BarChart(); bar.type = "col"; bar.title = None; bar.legend = None
        bar.width, bar.height = CW, CH
        bar.add_data(Reference(calc, min_col=2, min_row=1, max_row=ntop + 1),
                     titles_from_data=True)
        bar.set_categories(Reference(calc, min_col=1, min_row=2, max_row=ntop + 1))
        color_bar(bar, PALETTE[0])
        bar.y_axis.title = "Credits"
        line = LineChart()
        line.add_data(Reference(calc, min_col=3, min_row=1, max_row=ntop + 1),
                      titles_from_data=True)
        line.y_axis.axId = 200
        line.y_axis.title = "Cumulative %"
        line.y_axis.crosses = "max"
        if line.series:
            line.series[0].graphicalProperties = GraphicalProperties()
            line.series[0].graphicalProperties.line = LineProperties(solidFill="E15759", w=28000)
        bar += line
        dash.add_chart(bar, "B26")

    # 4. Credits by group (column)
    if ngrp:
        section("K25", "👥 Credits by group")
        ch = BarChart(); ch.type = "col"; ch.title = None; ch.legend = None
        ch.width, ch.height = CW, CH
        ch.add_data(Reference(calc, min_col=9, min_row=1, max_row=ngrp + 1),
                    titles_from_data=True)
        ch.set_categories(Reference(calc, min_col=8, min_row=2, max_row=ngrp + 1))
        color_bar(ch, PALETTE[3])
        dash.add_chart(ch, "K26")

    # 5. Utilization % by capped user (column)
    if nutil:
        section("B42", "📏 Utilization vs limit (%)")
        ch = BarChart(); ch.type = "col"; ch.title = None; ch.legend = None
        ch.width, ch.height = CW, CH
        ch.add_data(Reference(calc, min_col=12, min_row=1, max_row=nutil + 1),
                    titles_from_data=True)
        ch.set_categories(Reference(calc, min_col=11, min_row=2, max_row=nutil + 1))
        color_bar(ch, PALETTE[1])
        dash.add_chart(ch, "B43")

    # 6. Licence status (doughnut)
    if nlic:
        section("K42", "🪪 Licence status")
        ch = DoughnutChart(); ch.title = None
        ch.width, ch.height = CW, CH
        ch.add_data(Reference(calc, min_col=15, min_row=1, max_row=nlic + 1),
                    titles_from_data=True)
        ch.set_categories(Reference(calc, min_col=14, min_row=2, max_row=nlic + 1))
        color_points(ch, nlic)
        dash.add_chart(ch, "K43")

    # 7. Sessions vs credits (scatter)
    if nsc:
        section("B58", "🔬 Sessions vs credits")
        ch = ScatterChart(); ch.title = None; ch.legend = None
        ch.width, ch.height = CW, CH
        ch.x_axis.title = "Sessions"; ch.y_axis.title = "Credits"
        ch.x_axis.delete = False; ch.y_axis.delete = False
        xref = Reference(calc, min_col=17, min_row=2, max_row=nsc + 1)
        yref = Reference(calc, min_col=18, min_row=1, max_row=nsc + 1)
        s = Series(yref, xref, title_from_data=True)
        s.marker.symbol = "circle"; s.marker.size = 7
        s.graphicalProperties.line.noFill = True
        s.marker.graphicalProperties = GraphicalProperties(solidFill=PALETTE[4])
        ch.series.append(s)
        dash.add_chart(ch, "B59")

    # 8. Automation split (stacked column)
    if nauto:
        section("K58", "🤖 Automation split (scheduled vs manual)")
        ch = BarChart(); ch.type = "col"; ch.grouping = "stacked"; ch.overlap = 100
        ch.title = None; ch.width, ch.height = CW, CH
        ch.add_data(Reference(calc, min_col=21, max_col=22, min_row=1, max_row=nauto + 1),
                    titles_from_data=True)
        ch.set_categories(Reference(calc, min_col=20, min_row=2, max_row=nauto + 1))
        if len(ch.series) >= 2:
            ch.series[0].graphicalProperties = GraphicalProperties(solidFill=PALETTE[4])
            ch.series[1].graphicalProperties = GraphicalProperties(solidFill=PALETTE[1])
        dash.add_chart(ch, "K59")

    # --- Users sheet ---
    if ins["users"]:
        ws = wb.create_sheet("使用者 Users")
        cols = ["Display Name", "UPN", "Credit limit", "Credits used",
                "% Used (real)", "Sessions", "Credits/session", "Licensed",
                "Last active", "Flags"]
        ws.append(cols)
        style_header(ws, 1, len(cols))
        for u in sorted(ins["users"], key=lambda x: x["used"], reverse=True):
            flags = []
            if u["limit"] == 0 and u["used"] > 0:
                flags.append("Unlimited")
            if u["licensed"] and u["used"] == 0 and u["sessions"] == 0:
                flags.append("IdleLicence")
            if u["pct"] is not None and u["pct"] > OVER_LIMIT:
                flags.append("OverLimit")
            elif u["pct"] is not None and u["pct"] >= NEAR_LIMIT:
                flags.append("NearLimit")
            ws.append([
                u["name"], u["upn"],
                "無上限 Unlimited" if u["limit"] == 0 else int(u["limit"]),
                u["used"],
                "-" if u["pct"] is None else round(u["pct"] * 100, 1),
                u["sessions"],
                round(u["per_session"], 1),
                "Yes" if u["licensed"] else "No",
                u["last_active"],
                ", ".join(flags),
            ])
        autosize(ws, [22, 34, 16, 14, 14, 10, 14, 10, 14, 22])
        ws.freeze_panes = "A2"
        make_table(ws, "tblUsers")

    # --- Groups sheet ---
    if ins["groups"]:
        ws = wb.create_sheet("群組 Groups")
        cols = ["Group", "Total users", "Credits used", "Active members",
                "Avg/user/day", "Sessions", "Last active"]
        ws.append(cols)
        style_header(ws, 1, len(cols))
        for g in sorted(ins["groups"], key=lambda x: x["used"], reverse=True):
            ws.append([g["name"], int(g["total_users"]), g["used"],
                       int(g["active_members"]), round(g["avg_per_user_day"], 1),
                       int(g["sessions"]), g["last_active"]])
        autosize(ws, [28, 12, 14, 16, 14, 10, 14])
        ws.freeze_panes = "A2"
        make_table(ws, "tblGroups")

    # --- Services sheet + pie ---
    if ins["services"]:
        ws = wb.create_sheet("服務 Services")
        cols = ["Service", "Active users", "Credits used", "Last active"]
        ws.append(cols)
        style_header(ws, 1, len(cols))
        svc = sorted(ins["services"], key=lambda x: x["used"], reverse=True)
        for s in svc:
            ws.append([s["name"], int(s["active_users"]), s["used"], s["last_active"]])
        autosize(ws, [24, 14, 14, 14])
        make_table(ws, "tblServices")
        if len(svc) >= 1:
            pie = PieChart()
            pie.title = "Service split (credits)"
            pie.height = 8
            pie.width = 12
            data_ref = Reference(ws, min_col=3, min_row=1, max_row=len(svc) + 1)
            cats = Reference(ws, min_col=1, min_row=2, max_row=len(svc) + 1)
            pie.add_data(data_ref, titles_from_data=True)
            pie.set_categories(cats)
            ws.add_chart(pie, "F2")

    # --- Tasks sheet ---
    if ins["tasks"]:
        ws = wb.create_sheet("任務明細 Tasks")
        cols = ["Display Name", "UPN", "Total tasks", "Scheduled",
                "User-initiated", "Automation %", "Active days", "Last active"]
        ws.append(cols)
        style_header(ws, 1, len(cols))
        for t in sorted(ins["tasks"], key=lambda x: x["total"], reverse=True):
            ws.append([t["name"], t["upn"], int(t["total"]), int(t["scheduled"]),
                       int(t["user_initiated"]),
                       "-" if t["automation_ratio"] is None else round(t["automation_ratio"] * 100, 1),
                       int(t["active_days"]), t["last_active"]])
        autosize(ws, [22, 34, 12, 12, 14, 14, 12, 14])
        ws.freeze_panes = "A2"
        make_table(ws, "tblTasks")

    # --- Recommendations sheet ---
    ws = wb.create_sheet("最佳化建議 Recommendations")
    cols = ["優先級 Priority", "主題 Theme", "發現 Finding", "建議行動 Action", "預估效益 Benefit"]
    ws.append(cols)
    style_header(ws, 1, len(cols))
    if recs:
        for i, rec in enumerate(recs, start=2):
            ws.append([rec["priority"], rec["theme"], rec["finding"],
                       rec["action"], rec["benefit"]])
            fill = PRIO_FILL.get(rec["priority"])
            if fill:
                ws.cell(row=i, column=1).fill = fill
            for c in range(1, len(cols) + 1):
                ws.cell(row=i, column=c).alignment = Alignment(wrap_text=True, vertical="top")
    else:
        ws.append(["-", "無重大最佳化項目", "資料未觸發任何最佳化規則", "維持現狀並持續監控", "-"])
    autosize(ws, [14, 28, 46, 46, 28])
    ws.freeze_panes = "A2"
    make_table(ws, "tblRecs")

    wb.save(out_path)


# ---------------------------------------------------------------------------
# Word output
# ---------------------------------------------------------------------------
def write_word(ins: dict, recs: list[dict], out_path: str, meta: dict, found: dict) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    doc = Document()

    # CJK-friendly default font
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    def heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "Microsoft JhengHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        return h

    def kv_table(rows):
        t = doc.add_table(rows=0, cols=2)
        t.style = "Light Grid Accent 1"
        for k, v in rows:
            cells = t.add_row().cells
            cells[0].text = str(k)
            cells[1].text = str(v)
            for run in cells[0].paragraphs[0].runs:
                run.font.bold = True
        return t

    # Title
    title = doc.add_heading(f"AI FinOps 洞察報告", level=0)
    for run in title.runs:
        run.font.name = "Microsoft JhengHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    sub = doc.add_paragraph(f"{meta['org']}　|　{meta['period']}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Scope
    heading("報告範圍 Scope", 1)
    src_lines = []
    label = {"users": "使用者用量", "groups": "群組用量", "services": "服務用量", "tasks": "Cowork 任務明細"}
    for k, files in found.items():
        src_lines.append(f"{label.get(k, k)}（{', '.join(files)}）")
    kv_table([
        ("租戶 / 客戶", meta["org"]),
        ("報告期間", meta["period"]),
        ("產出時間", meta["generated"]),
        ("資料來源", "；".join(src_lines) if src_lines else "—"),
        ("計價單位", f"credits" + (f"（換算 {meta['currency']} @ {ins['credit_cost']}/credit）"
                                 if ins["credit_cost"] else "（僅以 credits 計，未換算貨幣）")),
    ])

    # Executive summary
    heading("執行摘要 Executive Summary", 1)
    cost_line = ""
    if ins["credit_cost"]:
        cost_line = f"，估算成本約 {ins['total_used'] * ins['credit_cost']:,.2f} {meta['currency']}"
    doc.add_paragraph(
        f"本期共消耗 {fmt_num(ins['total_used'])} credits{cost_line}，"
        f"活躍使用者 {ins['active_count']}/{ins['user_count']} 位，"
        f"平均每位活躍使用者 {fmt_num(ins['avg_per_active'])} credits。"
    )
    high = [r for r in recs if r["priority"] == "High"]
    if high:
        doc.add_paragraph("最優先的最佳化重點：")
        for r in high:
            doc.add_paragraph(f"{r['theme']}：{r['action']}", style="List Bullet")
    if ins["concentration"] >= CONCENTRATION_FLAG and ins["top"]:
        t = ins["top"][0]
        doc.add_paragraph(
            f"⚠️ 集中度偏高：單一使用者（{t['name'] or t['upn']}）即佔 "
            f"{fmt_pct(ins['concentration'])} 的總消耗。"
        )

    # Consumption
    heading("消耗總覽 Consumption Summary", 1)
    top = [u for u in ins["top"] if u["used"] > 0][:10]
    if top:
        doc.add_paragraph("消耗最高的使用者：")
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light List Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["使用者", "Credits", "佔比", "Sessions"]):
            hdr[i].text = h
        for u in top:
            cells = t.add_row().cells
            cells[0].text = u["name"] or u["upn"]
            cells[1].text = fmt_num(u["used"])
            cells[2].text = fmt_pct(u["used"] / ins["total_used"]) if ins["total_used"] else "-"
            cells[3].text = fmt_num(u["sessions"])
    if ins["services"]:
        doc.add_paragraph("依服務拆分：")
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light List Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["服務 Service", "Credits", "活躍使用者"]):
            hdr[i].text = h
        for s in sorted(ins["services"], key=lambda x: x["used"], reverse=True):
            cells = t.add_row().cells
            cells[0].text = s["name"]
            cells[1].text = fmt_num(s["used"])
            cells[2].text = fmt_num(s["active_users"])

    # Utilization
    heading("額度使用率 Utilization vs. Limits", 1)
    unlimited = [u for u in ins["users"] if u["limit"] == 0 and u["used"] > 0]
    capped = [u for u in ins["users"] if u["pct"] is not None]
    near = [u for u in capped if u["pct"] >= NEAR_LIMIT]
    doc.add_paragraph(
        f"無額度上限（unlimited）且有用量的帳號：{len(unlimited)} 位。"
        f"設有額度的帳號：{len(capped)} 位，其中 {len(near)} 位使用率 ≥ 80%。"
    )
    if unlimited:
        doc.add_paragraph(
            "無上限帳號是最大的成本治理缺口 — 這些帳號可無限消耗，建議優先設定月度額度。"
        )

    # Adoption
    if ins["tasks"]:
        heading("採用與參與 Adoption & Engagement", 1)
        tasks = [t for t in ins["tasks"] if t["automation_ratio"] is not None]
        if tasks:
            avg_auto = statistics.mean(t["automation_ratio"] for t in tasks)
            doc.add_paragraph(
                f"平均自動化比率（排程任務 ÷ 總任務）：{fmt_pct(avg_auto)}。"
                f"比率越高代表越能以排程取代臨時觸發，單位 credit 產出越高。"
            )

    # Anomalies
    heading("異常與風險 Anomalies & Risk", 1)
    wasted = [u for u in ins["users"] if u["licensed"] and u["used"] == 0 and u["sessions"] == 0]
    breaches = [u for u in ins["users"] if u["pct"] is not None and u["pct"] > OVER_LIMIT]
    anomaly_lines = []
    if ins["concentration"] >= CONCENTRATION_FLAG:
        anomaly_lines.append(f"集中度風險：最大使用者佔 {fmt_pct(ins['concentration'])}。")
    if wasted:
        anomaly_lines.append(f"閒置授權：{len(wasted)} 位持有授權但 0 用量。")
    if breaches:
        anomaly_lines.append(f"超出額度：{len(breaches)} 位使用者已超過月度上限。")
    if unlimited:
        anomaly_lines.append(f"未受管控：{len(unlimited)} 位無額度上限。")
    if anomaly_lines:
        for line in anomaly_lines:
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("本期未偵測到顯著異常。")

    # Optimization action plan (centerpiece)
    heading("成本最佳化行動計畫 Optimization Action Plan", 1)
    doc.add_paragraph("依優先級排序（High → Low）：")
    if recs:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(["優先級", "主題", "發現", "建議行動", "預估效益"]):
            hdr[i].text = h
            for run in hdr[i].paragraphs[0].runs:
                run.font.bold = True
        for r in recs:
            cells = t.add_row().cells
            cells[0].text = r["priority"]
            cells[1].text = r["theme"]
            cells[2].text = r["finding"]
            cells[3].text = r["action"]
            cells[4].text = r["benefit"]
    else:
        doc.add_paragraph("資料未觸發任何最佳化規則；維持現狀並持續每月監控。")

    # Method / disclaimer
    heading("附註與方法 Notes & Method", 1)
    snapshot = "本報告為單月快照；如需趨勢與預測，請保留每月匯出並逐月比較。"
    doc.add_paragraph(snapshot, style="List Bullet")
    doc.add_paragraph("所有金額以 credits 表示；除非提供實際換算率，否則不換算貨幣，避免誤導。",
                      style="List Bullet")
    doc.add_paragraph("limit = 0 代表『無上限』，非『零預算』。", style="List Bullet")

    doc.save(out_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> int:
    ap = argparse.ArgumentParser(description="AI FinOps report generator")
    ap.add_argument("--input-dir", required=True, help="Folder containing the CSV exports")
    ap.add_argument("--output-dir", required=True, help="Where to write the .docx / .xlsx")
    ap.add_argument("--period", default=datetime.now().strftime("%Y-%m"),
                    help="Reporting period label, e.g. 2026-07")
    ap.add_argument("--org", default="M365 Copilot Cowork", help="Tenant / customer name")
    ap.add_argument("--credit-cost", type=float, default=None,
                    help="Optional real credit->currency rate (never fabricate)")
    ap.add_argument("--currency", default="USD", help="Currency label for --credit-cost")
    args = ap.parse_args()

    if not os.path.isdir(args.input_dir):
        print(f"ERROR: input dir not found: {args.input_dir}", file=sys.stderr)
        return 2
    os.makedirs(args.output_dir, exist_ok=True)

    print(f"Scanning {args.input_dir} ...")
    data, found = load_csvs(args.input_dir)
    if not any(data.values()):
        print("ERROR: no recognizable Copilot consumption CSVs found.", file=sys.stderr)
        return 3
    for kind, files in found.items():
        print(f"  + {kind}: {', '.join(files)}")

    ins = compute_insights(data, args.credit_cost)
    recs = build_recommendations(ins)

    meta = {
        "org": args.org,
        "period": args.period,
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "currency": args.currency,
    }

    safe = args.period.replace("/", "-")
    xlsx_path = os.path.join(args.output_dir, f"AI-FinOps-Report_{safe}.xlsx")
    docx_path = os.path.join(args.output_dir, f"AI-FinOps-Report_{safe}.docx")

    write_excel(ins, recs, xlsx_path, meta)
    write_word(ins, recs, docx_path, meta, found)

    # Console summary
    print("\n=== FinOps summary ===")
    print(f"  Total credits used : {fmt_num(ins['total_used'])}")
    print(f"  Active users       : {ins['active_count']}/{ins['user_count']}")
    print(f"  Avg / active user  : {fmt_num(ins['avg_per_active'])}")
    print(f"  Top-user share     : {fmt_pct(ins['concentration'])}")
    print(f"  Recommendations    : {len(recs)}")
    for r in recs:
        print(f"    [{r['priority']}] {r['theme']}")
    print(f"\nWrote:\n  {xlsx_path}\n  {docx_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
